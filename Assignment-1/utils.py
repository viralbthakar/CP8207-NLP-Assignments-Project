import os
import docx
import string
import random
import zipfile
from collections import defaultdict
from nltk.tokenize import RegexpTokenizer


big_data_dict = defaultdict(list)


def styled_print(text, header=False):
    """Custom Print Function"""
    class style:
        BOLD = '\033[1m'
        UNDERLINE = '\033[4m'
        END = '\033[0m'

    if header:
        print(f'{style.BOLD}› {style.UNDERLINE}{text}{style.END}')
    else:
        print(f'    {text}')


def create_dir(root_dir, new_dir, header=True):
    styled_print(
        f'creating directory ... {os.path.join(root_dir, new_dir)}', header=header)
    os.makedirs(os.path.join(root_dir, new_dir), exist_ok=True)
    return os.path.join(root_dir, new_dir)


def random_select_dict(ip_dict, num_items):
    list_keys = random.choices(list(ip_dict.keys()), k=num_items)
    out_dict = {}
    for key in list_keys:
        out_dict[key] = ip_dict[key]
    return out_dict


def read_docx_file(file_path):
    zip_obj = zipfile.ZipFile(file_path)
    return zip_obj


def extract_images(file_path, out_path, extensions=[".jpg", ".jpeg", ".png", ".bmp"], verbose=0):
    styled_print(f"Extracting Images from {file_path}", header=True)
    image_file_paths = []
    zip_obj = read_docx_file(file_path)
    file_list = zip_obj.namelist()
    for file_name in file_list:
        _, extension = os.path.splitext(file_name)
        if extension in extensions:
            out_file_name = os.path.join(out_path, os.path.basename(file_name))
            if verbose:
                styled_print(f"Writing Image {file_name} to {out_file_name}")
            image_file_paths.append(out_file_name)
            with open(out_file_name, "wb") as out_file:
                out_file.write(zip_obj.read(file_name))
    return image_file_paths


def extract_paragraphs(file_path, out_path=None, min_char_count=1):
    styled_print(f"Extracting Paragraphs from {file_path}", header=True)
    paragraphs = {}
    document = docx.Document(file_path)
    for i in range(2, len(document.paragraphs)):
        if min_char_count is not None:
            if len(document.paragraphs[i].text) >= min_char_count:
                paragraphs[i] = document.paragraphs[i].text
        else:
            paragraphs[i] = document.paragraphs[i].text
    return paragraphs

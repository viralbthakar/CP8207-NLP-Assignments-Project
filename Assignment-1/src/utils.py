import os
import docx
import time
import string
import random
import zipfile
import seaborn as sns
import matplotlib.pyplot as plt
from collections import defaultdict
from nltk.tokenize import RegexpTokenizer


big_data_dict = defaultdict(list)


def styled_print(text, header=False):
    """Custom Print Function"""
    class style:
        BOLD = '\033[1m'
        UNDERLINE = '\033[4m'
        END = '\033[0m'

    if header:
        print(f'{style.BOLD}› {style.UNDERLINE}{text}{style.END}')
    else:
        print(f'    › {text}')
    time.sleep(0.5)


def create_dir(root_dir, new_dir, header=True):
    styled_print(
        f'creating directory ... {os.path.join(root_dir, new_dir)}', header=header)
    os.makedirs(os.path.join(root_dir, new_dir), exist_ok=True)
    return os.path.join(root_dir, new_dir)


def random_select_dict(ip_dict, num_items):
    list_keys = random.choices(list(ip_dict.keys()), k=num_items)
    out_dict = {}
    for key in list_keys:
        out_dict[key] = ip_dict[key]
    return out_dict


def read_docx_file(file_path):
    zip_obj = zipfile.ZipFile(file_path)
    return zip_obj


def extract_images(file_path, out_path, extensions=[".jpg", ".jpeg", ".png", ".bmp"], verbose=0):
    styled_print(f"Extracting Images from {file_path}", header=True)
    image_file_paths = []
    zip_obj = read_docx_file(file_path)
    file_list = zip_obj.namelist()
    for file_name in file_list:
        _, extension = os.path.splitext(file_name)
        if extension in extensions:
            out_file_name = os.path.join(out_path, os.path.basename(file_name))
            if verbose:
                styled_print(f"Writing Image {file_name} to {out_file_name}")
            image_file_paths.append(out_file_name)
            with open(out_file_name, "wb") as out_file:
                out_file.write(zip_obj.read(file_name))
    return image_file_paths


def extract_paragraphs(file_path, out_path=None, min_char_count=1):
    styled_print(f"Extracting Paragraphs from {file_path}", header=True)
    paragraphs = {}
    document = docx.Document(file_path)
    for i in range(2, len(document.paragraphs)):
        if min_char_count is not None:
            if len(document.paragraphs[i].text) >= min_char_count:
                paragraphs[i] = document.paragraphs[i].text
        else:
            paragraphs[i] = document.paragraphs[i].text
    return paragraphs


def plot_box_plot_hist_plot(df, column, title="Distribution Plot", figsize=(16, 16),
                            dpi=300, save_flag=False, file_path=None):
    fig, (ax_box, ax_hist) = plt.subplots(
        nrows=2,
        sharex=True,
        figsize=figsize,
        gridspec_kw={"height_ratios": (.20, .80)},
        dpi=dpi,
        constrained_layout=False
    )
    sns.boxplot(data=df, x=column, ax=ax_box)
    sns.histplot(data=df, x=column, ax=ax_hist, kde=True, bins='sqrt')
    ax_box.set(xlabel='')
    ax_box.set_facecolor('white')
    ax_hist.set_facecolor('white')
    plt.title(title)
    if save_flag:
        fig.savefig(file_path, dpi=dpi, facecolor='white')
        plt.close()


def plot_count_plot(df, column, hue=None, title="Count Plot", figsize=(24, 24), dpi=300,
                    save_flag=False, file_path=None):
    fig, axs = plt.subplots(1, 1, figsize=figsize,
                            dpi=dpi, constrained_layout=False)
    pt = sns.countplot(data=df, x=column, hue=hue,
                       palette=sns.color_palette("Set2"))
    pt.set_xticklabels(pt.get_xticklabels(), rotation=30)
    if hue is not None:
        axs.legend(loc="upper right", title=hue)
    axs.set_facecolor('white')
    plt.title(title)
    if save_flag:
        fig.savefig(file_path, dpi=dpi, facecolor='white')
        plt.close()


def combine_multiple_text_files(data_path):
    import glob
    if os.path.exists(os.path.join(data_path, "combined.txt")):
        os.remove(os.path.join(data_path, "combined.txt"))

    read_files = glob.glob(os.path.join(data_path, "*.txt"))
    with open(os.path.join(data_path, "combined.txt"), "wb") as outfile:
        for f in read_files:
            with open(f, "rb") as infile:
                outfile.write(infile.read())
    return os.path.join(data_path, "combined.txt")
